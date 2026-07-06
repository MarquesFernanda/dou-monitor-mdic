from __future__ import annotations

import logging
from datetime import date
from pathlib import Path
from typing import List, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .dou_client import Publicacao

logger = logging.getLogger(__name__)


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _style_header_cell(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shade_cell(cell, "1F4E79")


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Adiciona um hyperlink clicável de verdade dentro de um parágrafo do docx
    (a API pública do python-docx não tem esse recurso pronto, então
    manipulamos o XML interno do documento — é a forma padrão de fazer isso)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    new_run.append(rpr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_report(
    publicacoes: List[Publicacao],
    reference_date: date,
    output_path: Union[str, Path],
) -> Path:
    """Gera o .docx e devolve o Path do arquivo salvo."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    titulo = doc.add_heading("Monitoramento Diário DOU - MDIC/Inmetro", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(
        f"Referente à edição do DOU de {reference_date.strftime('%d/%m/%Y')} "
        f"— Seção 2 (órgãos: MDIC e Inmetro)"
    )
    run.italic = True

    doc.add_paragraph()

    if not publicacoes:
        doc.add_paragraph(
            f"Nenhuma publicação encontrada para o dia "
            f"{reference_date.strftime('%d/%m/%Y')}."
        )
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = [
            "Data da Publicação",
            "Órgão/Seção",
            "Título/Assunto",
            "Número do Ato",
            "Link para o DOU",
        ]
        for cell, header in zip(table.rows[0].cells, headers):
            _style_header_cell(cell, header)

        for pub in publicacoes:
            row = table.add_row().cells
            row[0].text = pub.data_publicacao
            row[1].text = pub.orgao
            row[2].text = pub.titulo
            row[3].text = pub.numero_ato
            row[4].paragraphs[0].text = ""  # limpa antes de inserir o hyperlink
            _add_hyperlink(row[4].paragraphs[0], pub.link, "Abrir no DOU")

        doc.add_paragraph()
        doc.add_paragraph(f"Total de publicações encontradas: {len(publicacoes)}")

    doc.add_paragraph()
    rodape = doc.add_paragraph()
    run = rodape.add_run(
        "Relatório gerado automaticamente pela automação de monitoramento do DOU."
    )
    run.font.size = Pt(8)
    run.italic = True

    doc.save(str(output_path))
    logger.info("Relatório salvo em %s", output_path)
    return output_path
